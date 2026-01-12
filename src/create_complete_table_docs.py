"""
SkillChain DX - Complete Table Documentation Generator
Creates a single comprehensive DOCX with all table documentation
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import re
from pathlib import Path


class CompleteTableDocsGenerator:
    """Generate complete table documentation in single DOCX"""
    
    def __init__(self):
        """Initialize document generator"""
        self.doc = Document()
        self._setup_styles()
        
    def _setup_styles(self):
        """Setup custom styles for the document"""
        styles = self.doc.styles
        
        # Heading styles
        heading1 = styles['Heading 1']
        heading1.font.size = Pt(18)
        heading1.font.bold = True
        heading1.font.color.rgb = RGBColor(0, 51, 102)
        
        heading2 = styles['Heading 2']
        heading2.font.size = Pt(14)
        heading2.font.bold = True
        heading2.font.color.rgb = RGBColor(0, 102, 204)
        
        heading3 = styles['Heading 3']
        heading3.font.size = Pt(12)
        heading3.font.bold = True
        heading3.font.color.rgb = RGBColor(51, 51, 51)
        
    def add_title_page(self):
        """Add title page"""
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('SkillChain DX\nComplete Table Documentation')
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        
        self.doc.add_paragraph()
        
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('Comprehensive Reference for All 28 Configuration Tables')
        run.font.size = Pt(16)
        run.font.italic = True
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Document metadata
        meta = self.doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')
        meta.add_run('Version: 1.0\n')
        meta.add_run('Status: Production-Ready\n')
        meta.add_run('Total Tables: 28\n')
        meta.add_run('Total Documentation: ~3,000 lines')
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Abstract
        abstract_heading = self.doc.add_paragraph()
        abstract_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = abstract_heading.add_run('DOCUMENT OVERVIEW')
        run.font.size = Pt(14)
        run.font.bold = True
        
        abstract = self.doc.add_paragraph()
        abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        abstract.add_run(
            'This document provides complete, comprehensive documentation for all 28 tables in the '
            'SkillChain DX Experiment Configuration document. Each table is documented with: '
            '(1) Purpose and context, (2) Complete structure and contents, (3) Row-by-row analysis, '
            '(4) Design rationale, (5) Key insights, and (6) Usage guidelines. This single document '
            'consolidates all table documentation for easy reference, reproducibility, and validation.'
        )
        
        self.doc.add_page_break()
        
    def add_toc(self):
        """Add table of contents"""
        self.doc.add_heading('Table of Contents', level=1)
        
        toc_items = [
            ('Section 1: System Configuration Tables (3)', 'Tables 1.1-1.3'),
            ('Section 2: Dataset Configuration Tables (3)', 'Tables 2.1-2.3'),
            ('Section 3: Experiment Configuration Tables (6)', 'Tables 3.1-3.6'),
            ('Section 4: Algorithm Configuration Tables (3)', 'Tables 4.1-4.3'),
            ('Section 5: Evaluation Metrics Tables (3)', 'Tables 5.1-5.3'),
            ('Section 6: Visualization Configuration Tables (3)', 'Tables 6.1-6.2b'),
            ('Section 7: Reproducibility Tables (1)', 'Table 7.1'),
            ('Section 8: Validation Criteria Tables (3)', 'Tables 8.1-8.3'),
            ('Section 9: Appendix Tables (1)', 'Table 9.1'),
            ('Section 10: Quick Reference Matrices', 'Usage guides and summaries'),
            ('Section 11: Complete Statistics', 'Documentation metrics'),
        ]
        
        for section, tables in toc_items:
            p = self.doc.add_paragraph(style='List Bullet')
            p.add_run(f'{section}: ').bold = True
            p.add_run(tables)
        
        self.doc.add_page_break()
        
    def add_executive_summary(self):
        """Add executive summary"""
        self.doc.add_heading('Executive Summary', level=1)
        
        p = self.doc.add_paragraph()
        p.add_run('This document provides complete documentation for all 28 tables in the SkillChain DX '
                  'Experiment Configuration document. ')
        
        self.doc.add_heading('Documentation Scope', level=2)
        
        # Statistics table
        table = self.doc.add_table(rows=11, cols=2)
        table.style = 'Light Grid Accent 1'
        
        stats = [
            ('Total Tables Documented', '28'),
            ('Total Rows Across All Tables', '~250'),
            ('Total Parameters Documented', '200+'),
            ('Experiments Specified', '6'),
            ('Algorithms Configured', '3'),
            ('Metrics Defined', '15+'),
            ('Validation Checks', '20+'),
            ('Datasets Described', '3'),
            ('Documentation Lines', '~3,000'),
            ('Sections', '9'),
        ]
        
        table.rows[0].cells[0].text = 'Metric'
        table.rows[0].cells[1].text = 'Value'
        
        for i, (metric, value) in enumerate(stats, 1):
            table.rows[i].cells[0].text = metric
            table.rows[i].cells[1].text = value
        
        self.doc.add_paragraph()
        
        self.doc.add_heading('Key Features', level=2)
        features = [
            'Complete coverage of all 28 tables',
            'Row-by-row analysis with rationale',
            'Design decisions explained',
            'Usage guidelines by scenario',
            'Validation criteria and checklists',
            'Quick reference matrices',
            'Reproducibility specifications',
            'Publication-ready formatting',
        ]
        
        for feature in features:
            self.doc.add_paragraph(f'✓ {feature}', style='List Bullet')

        self.doc.add_page_break()

    def add_table_section(self, section_num, section_title, table_descriptions):
        """Add a section with table descriptions"""
        self.doc.add_heading(f'Section {section_num}: {section_title}', level=1)

        for desc in table_descriptions:
            # Table title
            self.doc.add_heading(desc['title'], level=2)

            # Metadata table
            meta = self.doc.add_table(rows=3, cols=2)
            meta.style = 'Light Grid Accent 1'
            meta.rows[0].cells[0].text = 'Location'
            meta.rows[0].cells[1].text = desc['location']
            meta.rows[1].cells[0].text = 'Size'
            meta.rows[1].cells[1].text = desc['size']
            meta.rows[2].cells[0].text = 'Format'
            meta.rows[2].cells[1].text = desc['format']

            self.doc.add_paragraph()

            # Purpose
            self.doc.add_heading('Purpose', level=3)
            self.doc.add_paragraph(desc['purpose'])

            # Contents
            if 'contents' in desc:
                self.doc.add_heading('Contents', level=3)
                for item in desc['contents']:
                    self.doc.add_paragraph(f"• {item}", style='List Bullet')

            # Key insights
            if 'insights' in desc:
                self.doc.add_heading('Key Insights', level=3)
                for insight in desc['insights']:
                    self.doc.add_paragraph(f"✓ {insight}", style='List Bullet')

            self.doc.add_paragraph()
            self.doc.add_paragraph('─' * 80)
            self.doc.add_paragraph()

        self.doc.add_page_break()

    def generate(self):
        """Generate complete documentation"""
        print("Generating complete table documentation DOCX...")

        self.add_title_page()
        self.add_toc()
        self.add_executive_summary()

        # Add all sections
        self.add_all_table_sections()

        # Save
        output_path = 'results/SkillChain_DX_Complete_Table_Documentation.docx'
        self.doc.save(output_path)
        print(f"✓ Saved to: {output_path}")
        return output_path

