"""
SkillChain DX - Complete Experiment Configuration Document Generator
Generates a comprehensive DOCX with all 28 tables and detailed documentation
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime


class ConfigDocGenerator:
    """Generate complete experiment configuration documentation"""
    
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
        
    def _setup_styles(self):
        """Setup document styles"""
        styles = self.doc.styles
        
        h1 = styles['Heading 1']
        h1.font.size = Pt(18)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 51, 102)
        
        h2 = styles['Heading 2']
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0, 102, 153)
        
        h3 = styles['Heading 3']
        h3.font.size = Pt(12)
        h3.font.bold = True
        h3.font.color.rgb = RGBColor(51, 51, 51)
    
    def _add_table(self, headers, rows, style='Light Grid Accent 1'):
        """Add a formatted table"""
        table = self.doc.add_table(rows=len(rows)+1, cols=len(headers))
        table.style = style
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Headers
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Data rows
        for i, row in enumerate(rows):
            for j, value in enumerate(row):
                table.rows[i+1].cells[j].text = str(value)
        
        self.doc.add_paragraph()
        return table
    
    def add_title_page(self):
        """Add title page"""
        for _ in range(3):
            self.doc.add_paragraph()
        
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('SkillChain DX')
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('Complete Experiment Configuration Document')
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 102, 153)
        
        self.doc.add_paragraph()
        
        desc = self.doc.add_paragraph()
        desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = desc.add_run('Comprehensive Specification of All Experimental Parameters,\n'
                          'Settings, Algorithms, and Validation Criteria')
        run.font.size = Pt(14)
        run.font.italic = True
        
        for _ in range(4):
            self.doc.add_paragraph()
        
        # Metadata
        meta = self.doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run('─' * 50 + '\n')
        meta.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')
        meta.add_run('Version: 1.0\n')
        meta.add_run('Total Tables: 28\n')
        meta.add_run('Status: Production-Ready\n')
        meta.add_run('─' * 50)
        
        self.doc.add_page_break()
    
    def add_toc(self):
        """Add table of contents"""
        self.doc.add_heading('Table of Contents', level=1)
        
        sections = [
            ('1. System Configuration', '3 tables - Hardware, Software, Model specs'),
            ('2. Dataset Configuration', '3 tables - Roles, Courses, Employees'),
            ('3. Experiment Configurations', '6 tables - All experiment parameters'),
            ('4. Algorithm Configurations', '3 tables - Matching, Gap Analysis, Blockchain'),
            ('5. Evaluation Metrics', '3 tables - Quality, Performance, Scalability'),
            ('6. Visualization Configuration', '3 tables - Plotting parameters'),
            ('7. Reproducibility', '1 table - Random seeds'),
            ('8. Validation Criteria', '3 tables - Data, Model, Results validation'),
            ('9. Appendix', '1 table - Version history'),
        ]
        
        for section, desc in sections:
            p = self.doc.add_paragraph()
            run = p.add_run(section)
            run.bold = True
            p.add_run(f' — {desc}')
        
        self.doc.add_page_break()
    
    def add_section1_system(self):
        """Section 1: System Configuration"""
        self.doc.add_heading('1. SYSTEM CONFIGURATION', level=1)
        
        # Overview
        p = self.doc.add_paragraph()
        p.add_run('This section specifies all hardware, software, and model configurations '
                  'required to reproduce the SkillChain DX experiments.')
        
        # Table 1.1: Hardware
        self.doc.add_heading('Table 1.1: Hardware Environment Specifications', level=2)
        
        p = self.doc.add_paragraph()
        p.add_run('Purpose: ').bold = True
        p.add_run('Defines minimum and recommended hardware requirements for running experiments.')
        
        self._add_table(
            ['Component', 'Specification'],
            [
                ['Processor', 'Intel Core i5/i7 or AMD equivalent (2.0+ GHz, 4+ cores)'],
                ['RAM', '8 GB minimum, 16 GB recommended'],
                ['Storage', '10 GB available space (SSD recommended)'],
                ['GPU', 'Not required (CPU-only inference)'],
                ['Network', 'Internet connection (first-time model download)'],
            ]
        )
        
        p = self.doc.add_paragraph()
        p.add_run('Key Insights: ').bold = True
        p.add_run('Designed for standard hardware; no expensive GPU needed. '
                  '16GB RAM enables testing up to 1000 employees.')

        # Table 1.1 Explanation
        self.doc.add_heading('Table 1.1 Explanation', level=3)
        explanations = [
            ('Processor', 'Multi-core CPU (4+ cores) recommended for parallel batch processing. '
             'The embedding model runs on CPU, so clock speed matters more than GPU. '
             '2.0+ GHz ensures real-time inference performance.'),
            ('RAM', '8GB is the minimum to load the model (~500MB) plus data structures. '
             '16GB recommended for scalability tests with 1000+ employees where '
             'larger matrices are created in memory.'),
            ('Storage', '10GB accommodates: model cache (~500MB), datasets (~5MB), '
             'results and figures (~100MB), plus working space. '
             'SSD recommended for faster model loading (5s vs 15s on HDD).'),
            ('GPU', 'Not required because all-MiniLM-L6-v2 is optimized for CPU inference. '
             'GPU can accelerate batch processing but is optional for this scale.'),
            ('Network', 'Required only for first-time model download from HuggingFace. '
             'After caching, experiments run fully offline.'),
        ]
        for param, explanation in explanations:
            p = self.doc.add_paragraph()
            p.add_run(f'• {param}: ').bold = True
            p.add_run(explanation)

        self.doc.add_paragraph()

        # Table 1.2: Software
        self.doc.add_heading('Table 1.2: Software Environment Specifications', level=2)

        p = self.doc.add_paragraph()
        p.add_run('Purpose: ').bold = True
        p.add_run('Specifies exact software versions for reproducible environment setup.')

        self._add_table(
            ['Component', 'Version', 'Purpose'],
            [
                ['Python', '3.8+', 'Core programming language'],
                ['sentence-transformers', '2.2.2', 'Embedding model framework'],
                ['pandas', '2.1.4', 'Data manipulation and analysis'],
                ['numpy', '1.26.2', 'Numerical computations'],
                ['matplotlib', '3.8.2', 'Visualization library'],
                ['seaborn', '0.13.0', 'Statistical visualizations'],
                ['scikit-learn', '1.3.2', 'Similarity metrics'],
                ['python-docx', '1.1.0', 'Document generation'],
                ['hashlib', 'Standard lib', 'SHA-256 hashing'],
                ['json', 'Standard lib', 'Data serialization'],
            ]
        )

        p = self.doc.add_paragraph()
        p.add_run('Key Insights: ').bold = True
        p.add_run('All packages available via pip. Exact versions ensure identical results.')

        # Table 1.2 Explanation
        self.doc.add_heading('Table 1.2 Explanation', level=3)
        explanations = [
            ('sentence-transformers', 'Core framework providing pre-trained models. Version 2.2.2 '
             'includes all-MiniLM-L6-v2 with optimized inference. Provides encode() method for embeddings.'),
            ('pandas', 'Handles CSV loading and DataFrame operations. Used for reading datasets, '
             'filtering results, and generating reports. Version 2.1.4 has performance optimizations.'),
            ('numpy', 'Powers all numerical computations including cosine similarity calculations. '
             'Version 1.26.2 ensures compatibility with other packages.'),
            ('scikit-learn', 'Provides cosine_similarity() function from sklearn.metrics.pairwise. '
             'Optimized C implementation for fast pairwise similarity computation.'),
            ('matplotlib/seaborn', 'Generate all 8 experimental figures. Seaborn provides heatmaps '
             'and statistical plots. Colorblind-friendly palettes ensure accessibility.'),
        ]
        for param, explanation in explanations:
            p = self.doc.add_paragraph()
            p.add_run(f'• {param}: ').bold = True
            p.add_run(explanation)

        self.doc.add_paragraph()

        # Table 1.3: Model Configuration
        self.doc.add_heading('Table 1.3: Model Configuration Specifications', level=2)

        p = self.doc.add_paragraph()
        p.add_run('Purpose: ').bold = True
        p.add_run('Documents all parameters of the primary embedding model.')

        self._add_table(
            ['Parameter', 'Value'],
            [
                ['Model Name', 'sentence-transformers/all-MiniLM-L6-v2'],
                ['Architecture', 'BERT-based transformer (6 layers)'],
                ['Embedding Dimension', '384'],
                ['Max Sequence Length', '256 tokens'],
                ['Model Size', '22.7 MB'],
                ['Training Data', 'MS MARCO, NLI datasets (1B+ pairs)'],
                ['Inference Speed', '~50ms per batch (10 sequences)'],
            ]
        )

        p = self.doc.add_paragraph()
        p.add_run('Key Insights: ').bold = True
        p.add_run('22.7 MB model size enables edge deployment. '
                  '384 dimensions balance expressiveness and speed.')

        # Table 1.3 Explanation
        self.doc.add_heading('Table 1.3 Explanation', level=3)
        explanations = [
            ('Model Name', 'all-MiniLM-L6-v2 is a distilled BERT model trained specifically for '
             'semantic similarity. "all" means trained on diverse data, "MiniLM" is the architecture, '
             '"L6" means 6 transformer layers, "v2" is the second version with improvements.'),
            ('Embedding Dimension', '384-dimensional vectors capture semantic meaning. Higher dimensions '
             '(768, 1024) offer marginal accuracy gains but 2-3× slower. 384 is optimal for skill matching.'),
            ('Max Sequence Length', '256 tokens (~200 words) is sufficient for skill descriptions. '
             'Longer inputs are truncated. Most skill lists are under 50 tokens.'),
            ('Training Data', 'Pre-trained on 1B+ sentence pairs from MS MARCO (search queries) and '
             'NLI datasets (natural language inference). This makes it excellent at understanding '
             'semantic similarity between skill descriptions.'),
            ('Inference Speed', '~50ms for 10 sequences means real-time recommendations possible. '
             'Batch processing (32 sequences) further improves throughput.'),
        ]
        for param, explanation in explanations:
            p = self.doc.add_paragraph()
            p.add_run(f'• {param}: ').bold = True
            p.add_run(explanation)

        self.doc.add_page_break()

    def add_section2_datasets(self):
        """Section 2: Dataset Configuration"""
        self.doc.add_heading('2. DATASET CONFIGURATION', level=1)

        p = self.doc.add_paragraph()
        p.add_run('This section describes the structure and content of all input datasets.')

        # Table 2.1: Job Roles
        self.doc.add_heading('Table 2.1: Job Roles Dataset Specifications', level=2)

        p = self.doc.add_paragraph()
        p.add_run('Purpose: ').bold = True
        p.add_run('Defines structure and content of job roles dataset (target positions).')

        self._add_table(
            ['Attribute', 'Specification'],
            [
                ['Total Records', '15 job roles'],
                ['Fields', 'role_id, role_name, description, required_skills'],
                ['Skill Format', 'Comma-separated text strings'],
                ['Domains Covered', 'Data Science, Engineering, Management, Analytics'],
                ['Skill Count Range', '5-12 skills per role'],
            ]
        )

        p = self.doc.add_paragraph()
        p.add_run('Sample Roles: ').bold = True
        p.add_run('Data Scientist, Machine Learning Engineer, Data Analyst, '
                  'Business Intelligence Developer, Data Architect, etc.')

        # Table 2.1 Explanation
        self.doc.add_heading('Table 2.1 Explanation', level=3)
        explanations = [
            ('Total Records (15)', 'Chosen to provide sufficient variety across 4 domains while '
             'remaining manageable for detailed analysis. Each employee is compared against all 15 roles.'),
            ('Fields', 'role_id provides unique identifier; role_name is human-readable; '
             'description gives context for embedding generation; required_skills lists specific competencies.'),
            ('Skill Format', 'Comma-separated format (e.g., "Python, Machine Learning, SQL") enables '
             'easy parsing while remaining human-readable. Skills are concatenated for embedding.'),
            ('Domains Covered', 'Data Science, Engineering, Management, and Analytics represent '
             'typical organizational functions. Enables testing cross-functional career transitions.'),
            ('Skill Count Range', '5-12 skills per role mirrors real job descriptions. Fewer skills '
             'would be too generic; more would be too specific. This range ensures meaningful differentiation.'),
        ]
        for param, explanation in explanations:
            p = self.doc.add_paragraph()
            p.add_run(f'• {param}: ').bold = True
            p.add_run(explanation)

        self.doc.add_paragraph()

        # Table 2.2: Training Courses
        self.doc.add_heading('Table 2.2: Training Courses Dataset Specifications', level=2)

        p = self.doc.add_paragraph()
        p.add_run('Purpose: ').bold = True
        p.add_run('Specifies training courses dataset for skill gap remediation.')

        self._add_table(
            ['Attribute', 'Specification'],
            [
                ['Total Records', '25 training courses'],
                ['Fields', 'course_id, course_name, provider, skills_taught, duration, difficulty'],
                ['Skill Format', 'Comma-separated text strings'],
                ['Providers', 'Coursera, edX, Udacity, LinkedIn Learning, Pluralsight'],
                ['Duration Range', '4-40 hours'],
                ['Difficulty Levels', 'Beginner, Intermediate, Advanced'],
            ]
        )

        # Table 2.3: Employees
        self.doc.add_heading('Table 2.3: Employees Dataset Specifications', level=2)

        p = self.doc.add_paragraph()
        p.add_run('Purpose: ').bold = True
        p.add_run('Describes employee profiles dataset (workforce seeking recommendations).')

        self._add_table(
            ['Attribute', 'Specification'],
            [
                ['Total Records', '10 employee profiles'],
                ['Fields', 'employee_id, name, current_role, completed_courses, years_experience'],
                ['Course History', 'Comma-separated course IDs'],
                ['Experience Range', '2-10 years'],
                ['Roles Represented', 'Data Analyst, Engineer, Scientist, Manager'],
            ]
        )

        p = self.doc.add_paragraph()
        p.add_run('Key Insight: ').bold = True
        p.add_run('10 employees × 15 roles = 150 total comparisons for analysis.')

        # Table 2.2 & 2.3 Explanation
        self.doc.add_heading('Table 2.2 & 2.3 Explanations', level=3)

        p = self.doc.add_paragraph()
        p.add_run('Training Courses Dataset: ').bold = True
        p.add_run('25 courses provide sufficient variety for skill gap remediation. '
                  'Multiple providers (Coursera, edX, etc.) ensure real-world availability. '
                  'Duration range (4-40 hours) allows matching course intensity to gap size. '
                  'Difficulty levels enable appropriate progression paths.')

        p = self.doc.add_paragraph()
        p.add_run('Employees Dataset: ').bold = True
        p.add_run('10 employees provide statistical validity (n>30 comparisons per threshold). '
                  'Experience range (2-10 years) represents growth-focused workforce. '
                  'Course history enables tracking of completed training and progression. '
                  'Diverse starting roles test cross-functional recommendations.')

        self.doc.add_page_break()

    def add_section3_experiments(self):
        """Section 3: Experiment Configurations"""
        self.doc.add_heading('3. EXPERIMENT CONFIGURATIONS', level=1)

        p = self.doc.add_paragraph()
        p.add_run('This section provides complete specifications for all 6 experiments.')

        # Table 3.1: Threshold Analysis
        self.doc.add_heading('Table 3.1: Experiment 1 - Similarity Threshold Analysis', level=2)

        p = self.doc.add_paragraph()
        p.add_run('Purpose: ').bold = True
        p.add_run('Determines optimal similarity threshold for quality recommendations.')

        self._add_table(
            ['Parameter', 'Value/Range'],
            [
                ['Thresholds Tested', '50%, 60%, 70%, 75%, 80%, 85%, 90%'],
                ['Number of Thresholds', '7'],
                ['Employees Evaluated', '10 (all employees)'],
                ['Roles Evaluated', '15 (all roles)'],
                ['Total Comparisons', '150 (10 × 15)'],
                ['Metrics Collected', 'Avg recommendations, qualified employees, avg top score'],
                ['Expected Outcome', 'Identify optimal threshold (hypothesis: 70-75%)'],
                ['Total Evaluations', '1,050 (7 thresholds × 150 comparisons)'],
            ]
        )

        p = self.doc.add_paragraph()
        p.add_run('Rationale: ').bold = True
        p.add_run('Wide range from permissive (50%) to strict (90%). '
                  '5% increments in critical 70-90% zone for fine-grained analysis.')

        # Table 3.1 Explanation
        self.doc.add_heading('Table 3.1 Explanation', level=3)
        p = self.doc.add_paragraph()
        p.add_run('This experiment answers: "What similarity threshold produces the best recommendations?" ')
        p = self.doc.add_paragraph()
        p.add_run('• Why 7 thresholds? ').bold = True
        p.add_run('Covers full practical range. Below 50% indicates poor match; above 90% is rare.')
        p = self.doc.add_paragraph()
        p.add_run('• Why 5% increments in 70-90%? ').bold = True
        p.add_run('This is the critical decision zone where threshold choice most impacts outcomes.')
        p = self.doc.add_paragraph()
        p.add_run('• Expected outcome: ').bold = True
        p.add_run('70-75% threshold provides 3-5 recommendations per employee with >80% top scores.')

        self.doc.add_paragraph()

        # Table 3.2: Skill Gap Progression
        self.doc.add_heading('Table 3.2: Experiment 2 - Skill Gap Progression', level=2)

        p = self.doc.add_paragraph()
        p.add_run('Purpose: ').bold = True
        p.add_run('Tracks employee skill development through targeted training.')

        self._add_table(
            ['Parameter', 'Value/Range'],
            [
                ['Test Subject', 'Employee EMP001 (Data Analyst)'],
                ['Target Role', 'Data Strategy Officer'],
                ['Progression Stages', '4 (Initial, +1, +2, +3 courses)'],
                ['Courses Added', 'Data Governance, Leadership, Business Intelligence'],
                ['Metrics Tracked', 'Similarity score, skill gaps, courses completed'],
                ['Measurement Interval', 'After each course completion'],
                ['Expected Improvement', '+10-15% similarity score'],
                ['Expected Gap Reduction', '5 gaps → 1-2 gaps'],
                ['Hypothesis', 'Targeted courses yield measurable improvements'],
            ]
        )

        # Table 3.2 Explanation
        self.doc.add_heading('Table 3.2 Explanation', level=3)
        p = self.doc.add_paragraph()
        p.add_run('This experiment demonstrates the ROI of targeted training. ')
        p = self.doc.add_paragraph()
        p.add_run('• Why EMP001 as test subject? ').bold = True
        p.add_run('Data Analyst role is common starting point with clear upward paths.')
        p = self.doc.add_paragraph()
        p.add_run('• Why Data Strategy Officer target? ').bold = True
        p.add_run('Represents realistic 2-3 level promotion requiring new skills.')
        p = self.doc.add_paragraph()
        p.add_run('• Why 4 stages? ').bold = True
        p.add_run('Shows progression from initial state through incremental improvements.')
        p = self.doc.add_paragraph()
        p.add_run('• Expected +10-15% improvement: ').bold = True
        p.add_run('Each course should add 3-5% similarity by filling specific gaps.')

        self.doc.add_paragraph()

        # Table 3.3: Blockchain Performance
        self.doc.add_heading('Table 3.3: Experiment 3 - Blockchain Performance', level=2)

        p = self.doc.add_paragraph()
        p.add_run('Purpose: ').bold = True
        p.add_run('Evaluates blockchain credential system performance and scalability.')

        self._add_table(
            ['Parameter', 'Value/Range'],
            [
                ['Operations Tested', 'Issue credentials, Verify credentials'],
                ['Credential Counts', '10, 50, 100, 500'],
                ['Number of Scenarios', '8 (4 counts × 2 operations)'],
                ['Hash Algorithm', 'SHA-256'],
                ['Ledger Format', 'JSON (in-memory)'],
                ['Metrics Collected', 'Execution time (ms), throughput (ops/sec)'],
                ['Timing Method', 'Python time.time() with microsecond precision'],
                ['Repetitions', '1 per scenario (deterministic)'],
                ['Expected Throughput', '>1000 ops/sec for verification'],
                ['Hypothesis', 'Linear O(n) scalability'],
            ]
        )

        # Table 3.3 Explanation
        self.doc.add_heading('Table 3.3 Explanation', level=3)
        p = self.doc.add_paragraph()
        p.add_run('This experiment validates the blockchain credential system for production use. ')
        p = self.doc.add_paragraph()
        p.add_run('• Why SHA-256? ').bold = True
        p.add_run('Industry standard cryptographic hash. 256-bit output ensures collision resistance.')
        p = self.doc.add_paragraph()
        p.add_run('• Why these credential counts? ').bold = True
        p.add_run('10 (team), 50 (department), 100 (division), 500 (enterprise) represent real scales.')
        p = self.doc.add_paragraph()
        p.add_run('• Why >1000 ops/sec target? ').bold = True
        p.add_run('Enables real-time verification even at enterprise scale.')
        p = self.doc.add_paragraph()
        p.add_run('• Linear O(n) expectation: ').bold = True
        p.add_run('Each credential is independent; no interdependencies that would cause exponential growth.')

        self.doc.add_paragraph()

        # Table 3.4: Model Comparison
        self.doc.add_heading('Table 3.4: Experiment 4 - Embedding Model Comparison', level=2)

        p = self.doc.add_paragraph()
        p.add_run('Purpose: ').bold = True
        p.add_run('Compares three sentence-transformer models to validate model selection.')

        self._add_table(
            ['Parameter', 'Value/Range'],
            [
                ['Models Tested', 'all-MiniLM-L6-v2, all-MiniLM-L12-v2, paraphrase-MiniLM-L6-v2'],
                ['Number of Models', '3'],
                ['Test Dataset', '10 employees × 15 roles (150 pairs)'],
                ['Metrics Collected', 'Avg similarity, std dev, inference time, model size'],
                ['Model Sizes', 'L6: 22.7 MB, L12: 33.4 MB, Paraphrase: 22.7 MB'],
                ['Expected Winner', 'all-MiniLM-L6-v2 (speed/accuracy balance)'],
                ['Accuracy Tolerance', '±2% similarity score acceptable'],
                ['Hypothesis', 'Smaller models sufficient for skill matching'],
            ]
        )

        # Table 3.4 Explanation
        self.doc.add_heading('Table 3.4 Explanation', level=3)
        p = self.doc.add_paragraph()
        p.add_run('This experiment validates the choice of all-MiniLM-L6-v2 over alternatives. ')
        p = self.doc.add_paragraph()
        p.add_run('• Why these 3 models? ').bold = True
        p.add_run('L6 vs L12 tests layer depth impact; all vs paraphrase tests training data impact.')
        p = self.doc.add_paragraph()
        p.add_run('• Why ±2% tolerance? ').bold = True
        p.add_run('Differences under 2% are not practically significant for recommendations.')
        p = self.doc.add_paragraph()
        p.add_run('• Expected winner rationale: ').bold = True
        p.add_run('L6-v2 offers best speed/accuracy tradeoff. L12 is 47% larger with <2% gain.')

        self.doc.add_paragraph()

        # Table 3.5: Scalability Analysis
        self.doc.add_heading('Table 3.5: Experiment 5 - Scalability Analysis', level=2)

        p = self.doc.add_paragraph()
        p.add_run('Purpose: ').bold = True
        p.add_run('Evaluates system scalability across different organization sizes.')

        self._add_table(
            ['Parameter', 'Value/Range'],
            [
                ['Employee Counts', '10, 50, 100'],
                ['Role Counts', '15, 50, 100'],
                ['Scenarios Tested', '6 (selected combinations)'],
                ['Total Comparisons Range', '150 to 10,000'],
                ['Metrics Collected', 'Total time, time per comparison, memory usage'],
                ['Complexity Analysis', 'Linear regression (R² correlation)'],
                ['Expected Complexity', 'O(n×m) where n=employees, m=roles'],
                ['Expected R²', '>0.95 (strong linear correlation)'],
                ['Hypothesis', 'Linear scalability suitable for enterprise'],
            ]
        )

        # Table 3.5 Explanation
        self.doc.add_heading('Table 3.5 Explanation', level=3)
        p = self.doc.add_paragraph()
        p.add_run('This experiment proves the system can scale to enterprise organizations. ')
        p = self.doc.add_paragraph()
        p.add_run('• Why these scales? ').bold = True
        p.add_run('10×15=150 (startup), 50×50=2,500 (SMB), 100×100=10,000 (enterprise).')
        p = self.doc.add_paragraph()
        p.add_run('• Why R² > 0.95? ').bold = True
        p.add_run('Strong linear correlation means predictable performance at any scale.')
        p = self.doc.add_paragraph()
        p.add_run('• O(n×m) complexity: ').bold = True
        p.add_run('Each employee-role pair requires one comparison. No shortcuts possible for exhaustive matching.')
        p = self.doc.add_paragraph()
        p.add_run('• Extrapolation: ').bold = True
        p.add_run('1000 employees × 100 roles = 100,000 comparisons ≈ 10 seconds is acceptable.')

        self.doc.add_paragraph()

        # Table 3.6: Distribution Analysis
        self.doc.add_heading('Table 3.6: Experiment 6 - Score Distribution Analysis', level=2)

        p = self.doc.add_paragraph()
        p.add_run('Purpose: ').bold = True
        p.add_run('Analyzes statistical distribution of similarity scores.')

        self._add_table(
            ['Parameter', 'Value/Range'],
            [
                ['Sample Size', '150 (10 employees × 15 roles)'],
                ['Score Range', '0-100%'],
                ['Statistics Computed', 'Mean, median, std dev, min, max, quartiles'],
                ['Distribution Tests', 'Histogram, box plot, Q-Q plot'],
                ['Normality Assessment', 'Visual Q-Q plot analysis'],
                ['Per-Employee Analysis', 'Mean, max, min scores per employee'],
                ['Threshold Validation', 'Percentage exceeding 70%'],
                ['Expected Mean', '60-70%'],
                ['Expected Std Dev', '10-20%'],
                ['Hypothesis', 'Right-skewed distribution'],
            ]
        )

        # Table 3.6 Explanation
        self.doc.add_heading('Table 3.6 Explanation', level=3)
        p = self.doc.add_paragraph()
        p.add_run('This experiment characterizes the similarity score distribution. ')
        p = self.doc.add_paragraph()
        p.add_run('• Why these statistics? ').bold = True
        p.add_run('Mean, median, std dev, quartiles fully characterize the distribution shape.')
        p = self.doc.add_paragraph()
        p.add_run('• Why Q-Q plot? ').bold = True
        p.add_run('Visual test for normality. Deviations indicate skewness or heavy tails.')
        p = self.doc.add_paragraph()
        p.add_run('• Expected 60-70% mean: ').bold = True
        p.add_run('Skills are related but not identical across roles, so moderate similarity expected.')
        p = self.doc.add_paragraph()
        p.add_run('• Right-skewed hypothesis: ').bold = True
        p.add_run('Most comparisons moderate (50-70%), but some excellent matches (80%+) create right tail.')

        self.doc.add_page_break()

    def add_section4_algorithms(self):
        """Section 4: Algorithm Configurations"""
        self.doc.add_heading('4. ALGORITHM CONFIGURATIONS', level=1)

        p = self.doc.add_paragraph()
        p.add_run('This section documents all algorithm configurations and parameters.')

        # Table 4.1: Skill Matching
        self.doc.add_heading('Table 4.1: Skill Matching Algorithm Configuration', level=2)

        p = self.doc.add_paragraph()
        p.add_run('Purpose: ').bold = True
        p.add_run('Documents core skill matching algorithm that computes employee-role similarity.')

        self._add_table(
            ['Parameter', 'Configuration'],
            [
                ['Similarity Metric', 'Cosine similarity'],
                ['Embedding Method', 'Sentence-BERT (all-MiniLM-L6-v2)'],
                ['Text Preprocessing', 'Lowercase, whitespace normalization'],
                ['Skill Aggregation', 'Concatenate with space separator'],
                ['Normalization', 'L2 normalization (automatic)'],
                ['Batch Processing', 'Enabled (batch size: 32)'],
                ['Score Range', '0.0 to 1.0 (converted to 0-100%)'],
                ['Threshold Application', 'Post-computation filtering'],
                ['Ranking Method', 'Descending by similarity score'],
                ['Top-K Recommendations', 'K=5 (configurable)'],
            ]
        )

        p = self.doc.add_paragraph()
        p.add_run('Algorithm Flow: ').bold = True
        p.add_run('1. Concatenate skills → 2. Generate embeddings → '
                  '3. Compute cosine similarity → 4. Filter by threshold → 5. Rank results')

        # Table 4.1 Explanation
        self.doc.add_heading('Table 4.1 Explanation', level=3)
        p = self.doc.add_paragraph()
        p.add_run('• Cosine similarity: ').bold = True
        p.add_run('Measures angle between vectors, ignoring magnitude. Range [0,1] where 1 = identical direction.')
        p = self.doc.add_paragraph()
        p.add_run('• Skill aggregation: ').bold = True
        p.add_run('Concatenating skills ("Python Machine Learning SQL") captures semantic relationships.')
        p = self.doc.add_paragraph()
        p.add_run('• Batch size 32: ').bold = True
        p.add_run('Optimal for CPU cache utilization. Larger batches have diminishing returns.')
        p = self.doc.add_paragraph()
        p.add_run('• Top-K = 5: ').bold = True
        p.add_run('Shows best matches without overwhelming users. Configurable for different use cases.')

        self.doc.add_paragraph()

        # Table 4.2: Skill Gap Analysis
        self.doc.add_heading('Table 4.2: Skill Gap Analysis Algorithm Configuration', level=2)

        p = self.doc.add_paragraph()
        p.add_run('Purpose: ').bold = True
        p.add_run('Specifies algorithm for identifying skill gaps and recommending courses.')

        self._add_table(
            ['Parameter', 'Configuration'],
            [
                ['Gap Detection Method', 'Set difference (required - possessed)'],
                ['Skill Extraction', 'Parse comma-separated lists'],
                ['Matching Strategy', 'Exact string match (case-insensitive)'],
                ['Gap Prioritization', 'By frequency in high-similarity roles'],
                ['Course Recommendation', 'Match gap skills to course skills_taught'],
                ['Recommendation Limit', 'Top 3 courses per gap skill'],
                ['Course Ranking', 'By skill coverage and difficulty match'],
            ]
        )

        # Table 4.3: Blockchain Credential
        self.doc.add_heading('Table 4.3: Blockchain Credential Algorithm Configuration', level=2)

        p = self.doc.add_paragraph()
        p.add_run('Purpose: ').bold = True
        p.add_run('Documents blockchain-inspired credential verification algorithm.')

        self._add_table(
            ['Parameter', 'Configuration'],
            [
                ['Hash Function', 'SHA-256'],
                ['Hash Input', 'JSON string of credential data'],
                ['Credential ID Format', 'CRED_XXXX (sequential)'],
                ['Timestamp Format', 'ISO 8601 (YYYY-MM-DD HH:MM:SS)'],
                ['Previous Hash', 'SHA-256 hash of previous credential'],
                ['Genesis Block', 'Hash: 0000000000000000...'],
                ['Ledger Storage', 'In-memory list (production: database)'],
                ['Verification Method', 'Recompute hash and compare'],
                ['Immutability', 'Append-only ledger structure'],
            ]
        )

        p = self.doc.add_paragraph()
        p.add_run('Security: ').bold = True
        p.add_run('SHA-256 provides cryptographic integrity. Previous hash creates tamper-evident chain.')

        # Table 4.2 & 4.3 Explanations
        self.doc.add_heading('Table 4.2 & 4.3 Explanations', level=3)

        p = self.doc.add_paragraph()
        p.add_run('Skill Gap Analysis (4.2): ').bold = True
        p.add_run('Uses set difference to find missing skills. If role requires {A,B,C} and employee has {A,B}, '
                  'gap = {C}. Case-insensitive matching handles "Python" vs "python". '
                  'Top 3 courses per gap prevents overwhelming recommendations.')

        p = self.doc.add_paragraph()
        p.add_run('Blockchain Credentials (4.3): ').bold = True
        p.add_run('Each credential includes: employee_id, course_id, completion_date, issuer, plus hash of previous credential. '
                  'Genesis block uses zero hash. Verification recomputes hash and compares. '
                  'Any tampering breaks the chain, making fraud detectable.')

        self.doc.add_page_break()

    def add_section5_metrics(self):
        """Section 5: Evaluation Metrics"""
        self.doc.add_heading('5. EVALUATION METRICS', level=1)

        p = self.doc.add_paragraph()
        p.add_run('This section defines all metrics for evaluating system performance.')

        # Table 5.1: Quality Metrics
        self.doc.add_heading('Table 5.1: Recommendation Quality Metrics', level=2)

        p = self.doc.add_paragraph()
        p.add_run('Purpose: ').bold = True
        p.add_run('Defines metrics for evaluating recommendation quality and relevance.')

        self._add_table(
            ['Metric', 'Formula/Method', 'Target Value'],
            [
                ['Average Similarity Score', 'Mean of all scores', '60-70%'],
                ['Top Match Score', 'Highest score per employee', '>80%'],
                ['Recommendations per Employee', 'Count above threshold', '3-5'],
                ['Score Standard Deviation', 'Std dev of all scores', '10-20%'],
                ['Coverage', '% employees with ≥1 rec', '100%'],
                ['Precision@K', 'Relevant in top K', '>70%'],
            ]
        )

        # Table 5.2: Performance Metrics
        self.doc.add_heading('Table 5.2: Performance Metrics', level=2)

        p = self.doc.add_paragraph()
        p.add_run('Purpose: ').bold = True
        p.add_run('Specifies metrics for computational performance and efficiency.')

        self._add_table(
            ['Metric', 'Measurement Method', 'Target Value'],
            [
                ['Inference Time', 'time.time() before/after', '<100ms per batch'],
                ['Comparison Time', 'Time for cosine similarity', '<1ms per pair'],
                ['Total Processing Time', 'End-to-end for all comparisons', '<5s for 150 pairs'],
                ['Throughput', 'Operations per second', '>1000 ops/sec'],
                ['Memory Usage', 'Peak RAM during processing', '<500 MB'],
                ['Model Load Time', 'Time to load model', '<5 seconds'],
            ]
        )

        # Table 5.3: Scalability Metrics
        self.doc.add_heading('Table 5.3: Scalability Metrics', level=2)

        p = self.doc.add_paragraph()
        p.add_run('Purpose: ').bold = True
        p.add_run('Assesses system scalability to larger datasets.')

        self._add_table(
            ['Metric', 'Measurement Method', 'Target Value'],
            [
                ['Complexity Class', 'Big-O analysis', 'O(n×m)'],
                ['Linear Correlation (R²)', 'Regression of time vs comparisons', '>0.95'],
                ['Per-Comparison Time', 'Total time / comparisons', '50-100 μs'],
                ['Scalability Factor', 'Time ratio for 10× data', '<10×'],
                ['Extrapolated Performance', 'Projected time for 1000×100', '<10 seconds'],
            ]
        )

        # Table 5.1-5.3 Explanations
        self.doc.add_heading('Tables 5.1-5.3 Explanations', level=3)

        p = self.doc.add_paragraph()
        p.add_run('Quality Metrics (5.1): ').bold = True
        p.add_run('Average similarity 60-70% indicates meaningful differentiation. '
                  'Top match >80% ensures excellent opportunities exist. '
                  '3-5 recommendations per employee is manageable for review. '
                  '100% coverage ensures no employee is left without options.')

        p = self.doc.add_paragraph()
        p.add_run('Performance Metrics (5.2): ').bold = True
        p.add_run('<100ms inference enables real-time UI. '
                  '<5s total for 150 pairs is acceptable batch performance. '
                  '>1000 ops/sec for blockchain supports enterprise scale. '
                  '<500MB memory allows deployment on modest hardware.')

        p = self.doc.add_paragraph()
        p.add_run('Scalability Metrics (5.3): ').bold = True
        p.add_run('O(n×m) complexity is optimal for exhaustive comparison. '
                  'R² > 0.95 validates linear scaling assumption. '
                  '50-100 μs per comparison is achievable with batching. '
                  '<10 seconds for 100K comparisons enables enterprise deployment.')

        self.doc.add_page_break()

    def add_section6_visualization(self):
        """Section 6: Visualization Configuration"""
        self.doc.add_heading('6. VISUALIZATION CONFIGURATION', level=1)

        p = self.doc.add_paragraph()
        p.add_run('This section standardizes visualization parameters for consistency.')

        # Table 6.1: General Plotting
        self.doc.add_heading('Table 6.1: General Plotting Parameters', level=2)

        p = self.doc.add_paragraph()
        p.add_run('Purpose: ').bold = True
        p.add_run('Standardizes visualization parameters for publication quality.')

        self._add_table(
            ['Parameter', 'Value'],
            [
                ['Figure Size', '6 inches width (varies by type)'],
                ['DPI', '300 (publication quality)'],
                ['Font Family', 'DejaVu Sans'],
                ['Font Size', '10pt (labels), 12pt (titles)'],
                ['Color Palette', 'Seaborn default (colorblind-friendly)'],
                ['Grid Style', 'Light gray, alpha=0.3'],
                ['Legend Position', 'Best (automatic) or upper right'],
                ['File Format', 'PNG with tight layout'],
                ['Background', 'White'],
            ]
        )

        # Table 6.2a: Heatmap
        self.doc.add_heading('Table 6.2a: Heatmap Configuration', level=2)

        self._add_table(
            ['Parameter', 'Value'],
            [
                ['Colormap', 'YlOrRd (yellow-orange-red)'],
                ['Value Range', '0-100%'],
                ['Annotations', 'Enabled (score values in cells)'],
                ['Cell Size', 'Auto-scaled to fit'],
                ['Color Bar', 'Enabled with % labels'],
            ]
        )

        # Table 6.2b: Bar Chart
        self.doc.add_heading('Table 6.2b: Bar Chart Configuration', level=2)

        self._add_table(
            ['Parameter', 'Value'],
            [
                ['Orientation', 'Horizontal'],
                ['Bar Color', 'Steelblue (#4682B4)'],
                ['Value Labels', 'Enabled (end of bars)'],
                ['Sorting', 'By employee ID or score (descending)'],
            ]
        )

        # Table 6 Explanations
        self.doc.add_heading('Tables 6.1-6.2 Explanations', level=3)

        p = self.doc.add_paragraph()
        p.add_run('General Plotting (6.1): ').bold = True
        p.add_run('300 DPI ensures print-quality figures for publications. '
                  'DejaVu Sans is clean, professional, and widely available. '
                  'Colorblind-friendly palette ensures accessibility (affects ~8% of males). '
                  'Tight layout removes whitespace for efficient use of page space.')

        p = self.doc.add_paragraph()
        p.add_run('Heatmaps (6.2a): ').bold = True
        p.add_run('YlOrRd (yellow-orange-red) colormap is intuitive: yellow=low, red=high. '
                  'Works in grayscale for B&W printing. '
                  'Cell annotations provide exact values alongside color encoding.')

        p = self.doc.add_paragraph()
        p.add_run('Bar Charts (6.2b): ').bold = True
        p.add_run('Horizontal orientation allows longer labels (employee names, roles). '
                  'Steelblue is professional and high-contrast. '
                  'Value labels at bar ends eliminate need to read axis scales.')

        self.doc.add_page_break()

    def add_section7_reproducibility(self):
        """Section 7: Reproducibility"""
        self.doc.add_heading('7. REPRODUCIBILITY', level=1)

        p = self.doc.add_paragraph()
        p.add_run('This section ensures deterministic, reproducible results.')

        # Table 7.1: Random Seeds
        self.doc.add_heading('Table 7.1: Random Seed Configuration', level=2)

        p = self.doc.add_paragraph()
        p.add_run('Purpose: ').bold = True
        p.add_run('Documents random seed settings for deterministic results.')

        self._add_table(
            ['Component', 'Seed Value'],
            [
                ['NumPy Random', '42'],
                ['Python Random', '42'],
                ['Model Initialization', 'Deterministic (no randomness)'],
                ['Data Shuffling', 'Disabled (preserve order)'],
            ]
        )

        p = self.doc.add_paragraph()
        p.add_run('Reproducibility Guarantee: ').bold = True
        p.add_run('Same input → same embeddings → same similarity scores → '
                  'same recommendations. Result: Bit-for-bit reproducibility.')

        self.doc.add_heading('Reproducibility Checklist', level=2)

        checklist = [
            'Use exact software versions from Table 1.2',
            'Set random seeds before any processing',
            'Use same input data files (checksums verified)',
            'Follow experiment parameters exactly',
            'Validate results against Table 8.3 criteria',
        ]

        for item in checklist:
            self.doc.add_paragraph(f'☐ {item}', style='List Bullet')

        # Table 7.1 Explanation
        self.doc.add_heading('Table 7.1 Explanation', level=3)

        p = self.doc.add_paragraph()
        p.add_run('• Seed 42: ').bold = True
        p.add_run('Convention from "Hitchhiker\'s Guide to the Galaxy". Any fixed seed works.')

        p = self.doc.add_paragraph()
        p.add_run('• NumPy seed: ').bold = True
        p.add_run('Controls np.random.* functions used in array operations and sampling.')

        p = self.doc.add_paragraph()
        p.add_run('• Python random seed: ').bold = True
        p.add_run('Controls random.choice(), random.shuffle(), etc. in pure Python code.')

        p = self.doc.add_paragraph()
        p.add_run('• Model determinism: ').bold = True
        p.add_run('Pre-trained model has fixed weights. No random initialization during inference.')

        p = self.doc.add_paragraph()
        p.add_run('• Data order: ').bold = True
        p.add_run('Original CSV order preserved. No shuffling ensures same processing sequence.')

        self.doc.add_page_break()

    def add_section8_validation(self):
        """Section 8: Validation Criteria"""
        self.doc.add_heading('8. VALIDATION CRITERIA', level=1)

        p = self.doc.add_paragraph()
        p.add_run('This section provides checklists for validating data, model, and results.')

        # Table 8.1: Data Validation
        self.doc.add_heading('Table 8.1: Data Validation Criteria', level=2)

        p = self.doc.add_paragraph()
        p.add_run('Purpose: ').bold = True
        p.add_run('Checklist for validating input data quality.')

        self._add_table(
            ['Check', 'Criterion', 'Pass/Fail'],
            [
                ['Record Counts', 'Roles=15, Courses=25, Employees=10', 'Must match'],
                ['Missing Values', 'No null/empty required fields', 'Zero nulls'],
                ['Skill Format', 'Comma-separated strings', 'Valid format'],
                ['ID Uniqueness', 'All IDs unique within dataset', '100% unique'],
                ['Data Types', 'Correct types for all fields', 'All valid'],
                ['Encoding', 'UTF-8 encoding', 'UTF-8'],
            ]
        )

        # Table 8.2: Model Validation
        self.doc.add_heading('Table 8.2: Model Validation Criteria', level=2)

        p = self.doc.add_paragraph()
        p.add_run('Purpose: ').bold = True
        p.add_run('Checklist for validating embedding model functionality.')

        self._add_table(
            ['Check', 'Criterion', 'Pass/Fail'],
            [
                ['Model Loading', 'Successfully loads without errors', 'No errors'],
                ['Embedding Dimension', 'Output dimension = 384', '384'],
                ['Score Range', 'All scores in [0, 1]', '0 ≤ score ≤ 1'],
                ['Determinism', 'Same input → same output', '100% match'],
                ['Performance', 'Inference time < 100ms/batch', '<100ms'],
            ]
        )

        # Table 8.3: Results Validation
        self.doc.add_heading('Table 8.3: Results Validation Criteria', level=2)

        p = self.doc.add_paragraph()
        p.add_run('Purpose: ').bold = True
        p.add_run('Defines expected ranges for experimental results.')

        self._add_table(
            ['Check', 'Criterion', 'Expected Range'],
            [
                ['Mean Similarity', 'Average score across all pairs', '60-70%'],
                ['Std Deviation', 'Score variability', '10-20%'],
                ['Top Scores', 'Best match per employee', '>75%'],
                ['Coverage', 'Employees with recommendations', '100%'],
                ['Blockchain Throughput', 'Operations per second', '>1000 ops/sec'],
                ['Scalability R²', 'Linear correlation', '>0.95'],
                ['File Generation', 'All output files created', '12 files'],
            ]
        )

        # Table 8 Explanations
        self.doc.add_heading('Tables 8.1-8.3 Explanations', level=3)

        p = self.doc.add_paragraph()
        p.add_run('Data Validation (8.1): ').bold = True
        p.add_run('Run before experiments to catch data issues. '
                  'Record counts ensure complete data loaded. '
                  'Missing values would cause embedding errors. '
                  'ID uniqueness prevents duplicate processing. '
                  'UTF-8 encoding avoids character parsing issues.')

        p = self.doc.add_paragraph()
        p.add_run('Model Validation (8.2): ').bold = True
        p.add_run('Verifies model loaded correctly. '
                  '384 dimension confirms correct model variant. '
                  'Score range [0,1] ensures valid cosine similarity. '
                  'Determinism test: encode same text twice, compare outputs. '
                  'Performance test ensures no unexpected slowdowns.')

        p = self.doc.add_paragraph()
        p.add_run('Results Validation (8.3): ').bold = True
        p.add_run('Expected ranges based on experimental design. '
                  'Mean 60-70% reflects moderate cross-role similarity. '
                  'Std dev 10-20% shows meaningful differentiation. '
                  'Top scores >75% confirms quality matches exist. '
                  '100% coverage ensures no employee without recommendations.')

        self.doc.add_page_break()

    def add_section9_appendix(self):
        """Section 9: Appendix"""
        self.doc.add_heading('9. APPENDIX', level=1)

        # Table 9.1: Version History
        self.doc.add_heading('Table 9.1: Version History', level=2)

        self._add_table(
            ['Version', 'Date', 'Changes'],
            [
                ['1.0', '2026-01-12', 'Initial release with all 6 experiments'],
                ['1.1', 'TBD', 'Planned: Additional scalability tests'],
                ['2.0', 'TBD', 'Planned: Real-world deployment configuration'],
            ]
        )

        # Quick Reference
        self.doc.add_heading('Quick Reference: All 28 Tables', level=2)

        self._add_table(
            ['Section', 'Tables', 'Description'],
            [
                ['1. System Config', '1.1, 1.2, 1.3', 'Hardware, Software, Model'],
                ['2. Datasets', '2.1, 2.2, 2.3', 'Roles, Courses, Employees'],
                ['3. Experiments', '3.1-3.6', '6 experiment configurations'],
                ['4. Algorithms', '4.1, 4.2, 4.3', 'Matching, Gap, Blockchain'],
                ['5. Metrics', '5.1, 5.2, 5.3', 'Quality, Performance, Scale'],
                ['6. Visualization', '6.1, 6.2a, 6.2b', 'Plotting parameters'],
                ['7. Reproducibility', '7.1', 'Random seeds'],
                ['8. Validation', '8.1, 8.2, 8.3', 'Data, Model, Results'],
                ['9. Appendix', '9.1', 'Version history'],
            ]
        )

        # Statistics
        self.doc.add_heading('Documentation Statistics', level=2)

        self._add_table(
            ['Metric', 'Value'],
            [
                ['Total Tables', '28'],
                ['Total Rows', '~250'],
                ['Total Parameters', '200+'],
                ['Experiments Defined', '6'],
                ['Algorithms Configured', '3'],
                ['Metrics Specified', '15+'],
                ['Validation Checks', '20+'],
            ]
        )

        self.doc.add_paragraph()

        # Final note
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run('─' * 50 + '\n')
        p.add_run('End of Document\n').bold = True
        p.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')
        p.add_run('SkillChain DX - Complete Experiment Configuration\n')
        p.add_run('─' * 50)

    def generate(self):
        """Generate the complete document"""
        print("Generating complete experiment configuration document...")

        self.add_title_page()
        self.add_toc()
        self.add_section1_system()
        self.add_section2_datasets()
        self.add_section3_experiments()
        self.add_section4_algorithms()
        self.add_section5_metrics()
        self.add_section6_visualization()
        self.add_section7_reproducibility()
        self.add_section8_validation()
        self.add_section9_appendix()

        # Save document
        output_path = 'results/SkillChain_DX_Complete_Experiment_Configuration.docx'
        self.doc.save(output_path)
        print(f"✓ Document saved to: {output_path}")
        print(f"✓ Total sections: 9")
        print(f"✓ Total tables: 28")

        return output_path


if __name__ == '__main__':
    generator = ConfigDocGenerator()
    generator.generate()

