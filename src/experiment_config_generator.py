"""
SkillChain DX - Experiment Configuration Document Generator
Creates comprehensive DOCX documentation of all experimental configurations
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
from pathlib import Path
from datetime import datetime


class ExperimentConfigGenerator:
    """Generate comprehensive experiment configuration documentation"""
    
    def __init__(self):
        """Initialize document generator"""
        self.doc = Document()
        self._setup_styles()
        
    def _setup_styles(self):
        """Setup custom styles for the document"""
        styles = self.doc.styles
        
        # Heading styles
        heading1 = styles['Heading 1']
        heading1.font.size = Pt(16)
        heading1.font.bold = True
        heading1.font.color.rgb = RGBColor(0, 51, 102)
        
        heading2 = styles['Heading 2']
        heading2.font.size = Pt(14)
        heading2.font.bold = True
        heading2.font.color.rgb = RGBColor(0, 102, 204)
        
        heading3 = styles['Heading 3']
        heading3.font.size = Pt(12)
        heading3.font.bold = True
    
    def add_title_page(self):
        """Add title page"""
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('SkillChain DX\nExperimental Configuration Documentation')
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        
        self.doc.add_paragraph()
        
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('Complete Specification of Experimental Parameters,\nSettings, and Configurations')
        run.font.size = Pt(14)
        run.font.italic = True
        
        self.doc.add_paragraph()
        
        # Document metadata
        meta = self.doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')
        meta.add_run('Version: 1.0\n')
        meta.add_run('Status: Production-Ready')
        
        self.doc.add_paragraph()
        
        # Abstract
        abstract_heading = self.doc.add_paragraph()
        abstract_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = abstract_heading.add_run('DOCUMENT OVERVIEW')
        run.font.size = Pt(12)
        run.font.bold = True
        
        abstract = self.doc.add_paragraph()
        abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        abstract.add_run(
            'This document provides a comprehensive specification of all experimental configurations, '
            'parameters, and settings used in the SkillChain DX research project. It serves as a '
            'complete reference for reproducibility, validation, and future research extensions. '
            'All experiments, datasets, algorithms, and evaluation metrics are documented in detail '
            'with exact parameter values, rationale, and expected outcomes.'
        )
        
        self.doc.add_page_break()
    
    def add_system_configuration(self):
        """Add system configuration section"""
        self.doc.add_heading('1. SYSTEM CONFIGURATION', level=1)
        
        self.doc.add_heading('1.1 Hardware Environment', level=2)
        
        # Hardware specs table
        table = self.doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'
        
        table.rows[0].cells[0].text = 'Component'
        table.rows[0].cells[1].text = 'Specification'
        
        hw_data = [
            ('Processor', 'Intel Core i5/i7 or equivalent (2.0+ GHz)'),
            ('RAM', '8 GB minimum, 16 GB recommended'),
            ('Storage', '10 GB available space (SSD recommended)'),
            ('GPU', 'Not required (CPU-only inference)'),
            ('Network', 'Internet connection for model downloads')
        ]
        
        for i, (comp, spec) in enumerate(hw_data, 1):
            table.rows[i].cells[0].text = comp
            table.rows[i].cells[1].text = spec
        
        self.doc.add_paragraph()
        
        self.doc.add_heading('1.2 Software Environment', level=2)
        
        # Software specs table
        table = self.doc.add_table(rows=11, cols=3)
        table.style = 'Light Grid Accent 1'
        
        table.rows[0].cells[0].text = 'Component'
        table.rows[0].cells[1].text = 'Version'
        table.rows[0].cells[2].text = 'Purpose'
        
        sw_data = [
            ('Python', '3.8+', 'Core programming language'),
            ('sentence-transformers', '2.2.2', 'Embedding model framework'),
            ('pandas', '2.1.4', 'Data manipulation and analysis'),
            ('numpy', '1.26.2', 'Numerical computations'),
            ('matplotlib', '3.8.2', 'Visualization and plotting'),
            ('seaborn', '0.13.0', 'Statistical visualizations'),
            ('scikit-learn', '1.3.2', 'Similarity metrics'),
            ('python-docx', '1.1.0', 'Document generation'),
            ('hashlib', 'Standard library', 'SHA-256 hashing'),
            ('json', 'Standard library', 'Data serialization')
        ]
        
        for i, (comp, ver, purpose) in enumerate(sw_data, 1):
            table.rows[i].cells[0].text = comp
            table.rows[i].cells[1].text = ver
            table.rows[i].cells[2].text = purpose
        
        self.doc.add_paragraph()
        
        self.doc.add_heading('1.3 Model Configuration', level=2)
        
        self.doc.add_paragraph(
            'Primary Embedding Model: all-MiniLM-L6-v2'
        )

        # Model specs table
        table = self.doc.add_table(rows=8, cols=2)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Parameter'
        table.rows[0].cells[1].text = 'Value'

        model_data = [
            ('Model Name', 'sentence-transformers/all-MiniLM-L6-v2'),
            ('Architecture', 'BERT-based transformer (6 layers)'),
            ('Embedding Dimension', '384'),
            ('Max Sequence Length', '256 tokens'),
            ('Model Size', '22.7 MB'),
            ('Training Data', 'MS MARCO, NLI datasets (1B+ pairs)'),
            ('Inference Speed', '~50ms per batch (10 sequences)')
        ]

        for i, (param, value) in enumerate(model_data, 1):
            table.rows[i].cells[0].text = param
            table.rows[i].cells[1].text = value

        self.doc.add_paragraph()

    def add_dataset_configuration(self):
        """Add dataset configuration section"""
        self.doc.add_heading('2. DATASET CONFIGURATION', level=1)

        self.doc.add_heading('2.1 Job Roles Dataset', level=2)

        self.doc.add_paragraph(
            'File: data/job_roles.csv'
        )

        # Job roles table
        table = self.doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Attribute'
        table.rows[0].cells[1].text = 'Specification'

        jr_data = [
            ('Total Records', '15 job roles'),
            ('Fields', 'role_id, role_name, description, required_skills'),
            ('Skill Format', 'Comma-separated text strings'),
            ('Domains Covered', 'Data Science, Engineering, Management, Analytics'),
            ('Skill Count Range', '5-12 skills per role')
        ]

        for i, (attr, spec) in enumerate(jr_data, 1):
            table.rows[i].cells[0].text = attr
            table.rows[i].cells[1].text = spec

        self.doc.add_paragraph()

        self.doc.add_heading('2.2 Training Courses Dataset', level=2)

        self.doc.add_paragraph(
            'File: data/training_courses.csv'
        )

        # Courses table
        table = self.doc.add_table(rows=7, cols=2)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Attribute'
        table.rows[0].cells[1].text = 'Specification'

        tc_data = [
            ('Total Records', '25 training courses'),
            ('Fields', 'course_id, course_name, provider, skills_taught, duration, difficulty'),
            ('Skill Format', 'Comma-separated text strings'),
            ('Providers', 'Coursera, edX, Udacity, LinkedIn Learning, Pluralsight'),
            ('Duration Range', '4-40 hours'),
            ('Difficulty Levels', 'Beginner, Intermediate, Advanced')
        ]

        for i, (attr, spec) in enumerate(tc_data, 1):
            table.rows[i].cells[0].text = attr
            table.rows[i].cells[1].text = spec

        self.doc.add_paragraph()

        self.doc.add_heading('2.3 Employees Dataset', level=2)

        self.doc.add_paragraph(
            'File: data/employees.csv'
        )

        # Employees table
        table = self.doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Attribute'
        table.rows[0].cells[1].text = 'Specification'

        emp_data = [
            ('Total Records', '10 employee profiles'),
            ('Fields', 'employee_id, name, current_role, completed_courses, years_experience'),
            ('Course History', 'Comma-separated course IDs'),
            ('Experience Range', '2-10 years'),
            ('Roles Represented', 'Data Analyst, Engineer, Scientist, Manager')
        ]

        for i, (attr, spec) in enumerate(emp_data, 1):
            table.rows[i].cells[0].text = attr
            table.rows[i].cells[1].text = spec

        self.doc.add_paragraph()

    def add_experiment_configurations(self):
        """Add detailed experiment configurations"""
        self.doc.add_heading('3. EXPERIMENT CONFIGURATIONS', level=1)

        # Experiment 1
        self.doc.add_heading('3.1 Experiment 1: Similarity Threshold Analysis', level=2)

        self.doc.add_paragraph('Objective: Determine optimal similarity threshold for recommendations')

        table = self.doc.add_table(rows=9, cols=2)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Parameter'
        table.rows[0].cells[1].text = 'Value/Range'

        exp1_data = [
            ('Thresholds Tested', '50%, 60%, 70%, 75%, 80%, 85%, 90%'),
            ('Number of Thresholds', '7'),
            ('Employees Evaluated', '10 (all employees)'),
            ('Roles Evaluated', '15 (all roles)'),
            ('Total Comparisons', '150 (10 × 15)'),
            ('Metrics Collected', 'Avg recommendations, qualified employees, avg top score'),
            ('Expected Outcome', 'Identify threshold balancing quantity and quality'),
            ('Hypothesis', '70-75% threshold provides optimal balance')
        ]

        for i, (param, value) in enumerate(exp1_data, 1):
            table.rows[i].cells[0].text = param
            table.rows[i].cells[1].text = value

        self.doc.add_paragraph()

        # Experiment 2
        self.doc.add_heading('3.2 Experiment 2: Skill Gap Progression', level=2)

        self.doc.add_paragraph('Objective: Track skill development through targeted training')

        table = self.doc.add_table(rows=10, cols=2)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Parameter'
        table.rows[0].cells[1].text = 'Value/Range'

        exp2_data = [
            ('Test Subject', 'Employee EMP001 (Data Analyst)'),
            ('Target Role', 'Data Strategy Officer'),
            ('Progression Stages', '4 (Initial, +1 course, +2 courses, +3 courses)'),
            ('Courses Added', 'Data Governance, Leadership, Business Intelligence'),
            ('Metrics Tracked', 'Similarity score, skill gaps, courses completed'),
            ('Measurement Interval', 'After each course completion'),
            ('Expected Improvement', '+10-15% similarity score'),
            ('Expected Gap Reduction', '5 gaps → 1-2 gaps'),
            ('Hypothesis', 'Targeted courses yield measurable improvements')
        ]

        for i, (param, value) in enumerate(exp2_data, 1):
            table.rows[i].cells[0].text = param
            table.rows[i].cells[1].text = value

        self.doc.add_paragraph()

        # Experiment 3
        self.doc.add_heading('3.3 Experiment 3: Blockchain Performance', level=2)

        self.doc.add_paragraph('Objective: Evaluate blockchain credential operation performance')

        table = self.doc.add_table(rows=11, cols=2)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Parameter'
        table.rows[0].cells[1].text = 'Value/Range'

        exp3_data = [
            ('Operations Tested', 'Issue credentials, Verify credentials'),
            ('Credential Counts', '10, 50, 100, 500'),
            ('Number of Scenarios', '8 (4 counts × 2 operations)'),
            ('Hash Algorithm', 'SHA-256'),
            ('Ledger Format', 'JSON'),
            ('Metrics Collected', 'Execution time (ms), throughput (ops/sec)'),
            ('Timing Method', 'Python time.time() with microsecond precision'),
            ('Repetitions', '1 per scenario (deterministic operations)'),
            ('Expected Throughput', '>1000 ops/sec for verification'),
            ('Hypothesis', 'Linear O(n) scalability with credential count')
        ]

        for i, (param, value) in enumerate(exp3_data, 1):
            table.rows[i].cells[0].text = param
            table.rows[i].cells[1].text = value

        self.doc.add_paragraph()

        # Experiment 4
        self.doc.add_heading('3.4 Experiment 4: Embedding Model Comparison', level=2)

        self.doc.add_paragraph('Objective: Compare sentence-transformer models for optimal performance')

        table = self.doc.add_table(rows=10, cols=2)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Parameter'
        table.rows[0].cells[1].text = 'Value/Range'

        exp4_data = [
            ('Models Tested', 'all-MiniLM-L6-v2, all-MiniLM-L12-v2, paraphrase-MiniLM-L6-v2'),
            ('Number of Models', '3'),
            ('Test Dataset', 'Same 10 employees × 15 roles'),
            ('Metrics Collected', 'Avg similarity, std dev, inference time, model size'),
            ('Inference Timing', 'Average over all 150 comparisons'),
            ('Model Sizes', 'L6: 22.7 MB, L12: 33.4 MB, Paraphrase: 22.7 MB'),
            ('Expected Winner', 'all-MiniLM-L6-v2 (speed/accuracy balance)'),
            ('Accuracy Tolerance', '±2% similarity score acceptable'),
            ('Hypothesis', 'Smaller models sufficient for skill matching task')
        ]

        for i, (param, value) in enumerate(exp4_data, 1):
            table.rows[i].cells[0].text = param
            table.rows[i].cells[1].text = value

        self.doc.add_paragraph()

        # Experiment 5
        self.doc.add_heading('3.5 Experiment 5: Scalability Analysis', level=2)

        self.doc.add_paragraph('Objective: Evaluate system scalability across organization sizes')

        table = self.doc.add_table(rows=10, cols=2)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Parameter'
        table.rows[0].cells[1].text = 'Value/Range'

        exp5_data = [
            ('Employee Counts', '10, 50, 100'),
            ('Role Counts', '15, 50, 100'),
            ('Scenarios Tested', '6 (combinations of employee/role counts)'),
            ('Total Comparisons Range', '150 to 10,000'),
            ('Metrics Collected', 'Total time, time per comparison, memory usage'),
            ('Complexity Analysis', 'Linear regression (R² correlation)'),
            ('Expected Complexity', 'O(n×m) where n=employees, m=roles'),
            ('Expected R²', '>0.95 (strong linear correlation)'),
            ('Hypothesis', 'Linear scalability suitable for enterprise deployment')
        ]

        for i, (param, value) in enumerate(exp5_data, 1):
            table.rows[i].cells[0].text = param
            table.rows[i].cells[1].text = value

        self.doc.add_paragraph()

        # Experiment 6
        self.doc.add_heading('3.6 Experiment 6: Score Distribution Analysis', level=2)

        self.doc.add_paragraph('Objective: Analyze statistical distribution of similarity scores')

        table = self.doc.add_table(rows=11, cols=2)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Parameter'
        table.rows[0].cells[1].text = 'Value/Range'

        exp6_data = [
            ('Sample Size', '150 (10 employees × 15 roles)'),
            ('Score Range', '0-100%'),
            ('Statistics Computed', 'Mean, median, std dev, min, max, quartiles'),
            ('Distribution Tests', 'Histogram, box plot, Q-Q plot'),
            ('Normality Assessment', 'Visual Q-Q plot analysis'),
            ('Per-Employee Analysis', 'Mean, max, min scores for each employee'),
            ('Threshold Validation', 'Percentage of scores exceeding 70%'),
            ('Expected Mean', '60-70%'),
            ('Expected Std Dev', '10-20%'),
            ('Hypothesis', 'Right-skewed distribution with meaningful differentiation')
        ]

        for i, (param, value) in enumerate(exp6_data, 1):
            table.rows[i].cells[0].text = param
            table.rows[i].cells[1].text = value

        self.doc.add_paragraph()

    def add_algorithm_configurations(self):
        """Add algorithm and processing configurations"""
        self.doc.add_heading('4. ALGORITHM CONFIGURATIONS', level=1)

        self.doc.add_heading('4.1 Skill Matching Algorithm', level=2)

        table = self.doc.add_table(rows=11, cols=2)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Parameter'
        table.rows[0].cells[1].text = 'Configuration'

        algo_data = [
            ('Similarity Metric', 'Cosine similarity'),
            ('Embedding Method', 'Sentence-BERT (all-MiniLM-L6-v2)'),
            ('Text Preprocessing', 'Lowercase, whitespace normalization'),
            ('Skill Aggregation', 'Concatenate all skills with space separator'),
            ('Normalization', 'L2 normalization (automatic in model)'),
            ('Batch Processing', 'Enabled (batch size: 32)'),
            ('Score Range', '0.0 to 1.0 (converted to 0-100%)'),
            ('Threshold Application', 'Post-computation filtering'),
            ('Ranking Method', 'Descending by similarity score'),
            ('Top-K Recommendations', 'K=5 (configurable)')
        ]

        for i, (param, config) in enumerate(algo_data, 1):
            table.rows[i].cells[0].text = param
            table.rows[i].cells[1].text = config

        self.doc.add_paragraph()

        self.doc.add_heading('4.2 Skill Gap Analysis Algorithm', level=2)

        table = self.doc.add_table(rows=8, cols=2)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Parameter'
        table.rows[0].cells[1].text = 'Configuration'

        gap_data = [
            ('Gap Detection Method', 'Set difference (required - possessed)'),
            ('Skill Extraction', 'Parse comma-separated skill lists'),
            ('Matching Strategy', 'Exact string match (case-insensitive)'),
            ('Gap Prioritization', 'By frequency in high-similarity roles'),
            ('Course Recommendation', 'Match gap skills to course skills_taught'),
            ('Recommendation Limit', 'Top 3 courses per gap skill'),
            ('Course Ranking', 'By skill coverage and difficulty match')
        ]

        for i, (param, config) in enumerate(gap_data, 1):
            table.rows[i].cells[0].text = param
            table.rows[i].cells[1].text = config

        self.doc.add_paragraph()

        self.doc.add_heading('4.3 Blockchain Credential Algorithm', level=2)

        table = self.doc.add_table(rows=10, cols=2)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Parameter'
        table.rows[0].cells[1].text = 'Configuration'

        bc_data = [
            ('Hash Function', 'SHA-256'),
            ('Hash Input', 'JSON string of credential data'),
            ('Credential ID Format', 'CRED_XXXX (sequential)'),
            ('Timestamp Format', 'ISO 8601 (YYYY-MM-DD HH:MM:SS)'),
            ('Previous Hash', 'SHA-256 hash of previous credential'),
            ('Genesis Block', 'Hash: 0000000000000000...'),
            ('Ledger Storage', 'In-memory list (production: database)'),
            ('Verification Method', 'Recompute hash and compare'),
            ('Immutability', 'Append-only ledger structure')
        ]

        for i, (param, config) in enumerate(bc_data, 1):
            table.rows[i].cells[0].text = param
            table.rows[i].cells[1].text = config

        self.doc.add_paragraph()

    def add_evaluation_metrics(self):
        """Add evaluation metrics configuration"""
        self.doc.add_heading('5. EVALUATION METRICS', level=1)

        self.doc.add_heading('5.1 Recommendation Quality Metrics', level=2)

        table = self.doc.add_table(rows=7, cols=3)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Metric'
        table.rows[0].cells[1].text = 'Formula/Method'
        table.rows[0].cells[2].text = 'Target Value'

        quality_data = [
            ('Average Similarity Score', 'Mean of all employee-role scores', '60-70%'),
            ('Top Match Score', 'Highest score per employee', '>80%'),
            ('Recommendations per Employee', 'Count of scores above threshold', '3-5'),
            ('Score Standard Deviation', 'Std dev of all scores', '10-20%'),
            ('Coverage', '% of employees with ≥1 recommendation', '100%'),
            ('Precision@K', 'Relevant recommendations in top K', '>70%')
        ]

        for i, (metric, formula, target) in enumerate(quality_data, 1):
            table.rows[i].cells[0].text = metric
            table.rows[i].cells[1].text = formula
            table.rows[i].cells[2].text = target

        self.doc.add_paragraph()

        self.doc.add_heading('5.2 Performance Metrics', level=2)

        table = self.doc.add_table(rows=7, cols=3)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Metric'
        table.rows[0].cells[1].text = 'Measurement Method'
        table.rows[0].cells[2].text = 'Target Value'

        perf_data = [
            ('Inference Time', 'time.time() before/after embedding', '<100ms per batch'),
            ('Comparison Time', 'Time for cosine similarity computation', '<1ms per pair'),
            ('Total Processing Time', 'End-to-end for all comparisons', '<5s for 150 pairs'),
            ('Throughput', 'Operations per second', '>1000 ops/sec'),
            ('Memory Usage', 'Peak RAM during processing', '<500 MB'),
            ('Model Load Time', 'Time to load embedding model', '<5 seconds')
        ]

        for i, (metric, method, target) in enumerate(perf_data, 1):
            table.rows[i].cells[0].text = metric
            table.rows[i].cells[1].text = method
            table.rows[i].cells[2].text = target

        self.doc.add_paragraph()

        self.doc.add_heading('5.3 Scalability Metrics', level=2)

        table = self.doc.add_table(rows=6, cols=3)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Metric'
        table.rows[0].cells[1].text = 'Measurement Method'
        table.rows[0].cells[2].text = 'Target Value'

        scale_data = [
            ('Complexity Class', 'Big-O analysis of runtime', 'O(n×m)'),
            ('Linear Correlation (R²)', 'Regression of time vs comparisons', '>0.95'),
            ('Per-Comparison Time', 'Total time / number of comparisons', '50-100 μs'),
            ('Scalability Factor', 'Time ratio for 10× data increase', '<10×'),
            ('Extrapolated Performance', 'Projected time for 1000×100', '<10 seconds')
        ]

        for i, (metric, method, target) in enumerate(scale_data, 1):
            table.rows[i].cells[0].text = metric
            table.rows[i].cells[1].text = method
            table.rows[i].cells[2].text = target

        self.doc.add_paragraph()

    def add_visualization_configurations(self):
        """Add visualization configurations"""
        self.doc.add_heading('6. VISUALIZATION CONFIGURATIONS', level=1)

        self.doc.add_heading('6.1 General Plotting Parameters', level=2)

        table = self.doc.add_table(rows=10, cols=2)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Parameter'
        table.rows[0].cells[1].text = 'Value'

        viz_data = [
            ('Figure Size', '6 inches width (varies by plot type)'),
            ('DPI', '300 (publication quality)'),
            ('Font Family', 'DejaVu Sans'),
            ('Font Size', '10pt (labels), 12pt (titles)'),
            ('Color Palette', 'Seaborn default (colorblind-friendly)'),
            ('Grid Style', 'Light gray, alpha=0.3'),
            ('Legend Position', 'Best (automatic) or upper right'),
            ('File Format', 'PNG with tight layout'),
            ('Background', 'White')
        ]

        for i, (param, value) in enumerate(viz_data, 1):
            table.rows[i].cells[0].text = param
            table.rows[i].cells[1].text = value

        self.doc.add_paragraph()

        self.doc.add_heading('6.2 Plot-Specific Configurations', level=2)

        # Heatmap config
        self.doc.add_paragraph('Heatmap Configuration:')
        table = self.doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Parameter'
        table.rows[0].cells[1].text = 'Value'

        hm_data = [
            ('Colormap', 'YlOrRd (yellow-orange-red)'),
            ('Value Range', '0-100%'),
            ('Annotations', 'Enabled (score values in cells)'),
            ('Cell Size', 'Auto-scaled to fit'),
            ('Color Bar', 'Enabled with % labels')
        ]

        for i, (param, value) in enumerate(hm_data, 1):
            table.rows[i].cells[0].text = param
            table.rows[i].cells[1].text = value

        self.doc.add_paragraph()

        # Bar chart config
        self.doc.add_paragraph('Bar Chart Configuration:')
        table = self.doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Parameter'
        table.rows[0].cells[1].text = 'Value'

        bar_data = [
            ('Orientation', 'Horizontal'),
            ('Bar Color', 'Steelblue (#4682B4)'),
            ('Value Labels', 'Enabled (end of bars)'),
            ('Sorting', 'By employee ID or score (descending)')
        ]

        for i, (param, value) in enumerate(bar_data, 1):
            table.rows[i].cells[0].text = param
            table.rows[i].cells[1].text = value

        self.doc.add_paragraph()

    def add_reproducibility_guidelines(self):
        """Add reproducibility guidelines"""
        self.doc.add_heading('7. REPRODUCIBILITY GUIDELINES', level=1)

        self.doc.add_heading('7.1 Environment Setup', level=2)

        self.doc.add_paragraph('Step 1: Install Python 3.8 or higher')
        self.doc.add_paragraph('Step 2: Install required packages:')

        code = self.doc.add_paragraph()
        code.add_run('pip install -r requirements.txt').font.name = 'Courier New'

        self.doc.add_paragraph()
        self.doc.add_paragraph('Step 3: Verify installation:')

        code = self.doc.add_paragraph()
        code.add_run('python -c "import sentence_transformers; print(sentence_transformers.__version__)"').font.name = 'Courier New'

        self.doc.add_paragraph()

        self.doc.add_heading('7.2 Data Preparation', level=2)

        self.doc.add_paragraph('Step 1: Ensure data files are in data/ directory:')
        steps = [
            'data/job_roles.csv',
            'data/training_courses.csv',
            'data/employees.csv'
        ]
        for step in steps:
            p = self.doc.add_paragraph(step, style='List Bullet')

        self.doc.add_paragraph()
        self.doc.add_paragraph('Step 2: Verify data format matches specifications in Section 2')
        self.doc.add_paragraph('Step 3: Check for missing values or formatting errors')

        self.doc.add_paragraph()

        self.doc.add_heading('7.3 Experiment Execution', level=2)

        self.doc.add_paragraph('To reproduce all experiments:')

        code = self.doc.add_paragraph()
        code.add_run('python run_comprehensive_experiments.py').font.name = 'Courier New'

        self.doc.add_paragraph()
        self.doc.add_paragraph('To run individual experiments:')

        code = self.doc.add_paragraph()
        code.add_run('python src/experiments.py --experiment 1').font.name = 'Courier New'

        self.doc.add_paragraph()

        self.doc.add_heading('7.4 Random Seed Configuration', level=2)

        table = self.doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Component'
        table.rows[0].cells[1].text = 'Seed Value'

        seed_data = [
            ('NumPy Random', '42'),
            ('Python Random', '42'),
            ('Model Initialization', 'Deterministic (no randomness)'),
            ('Data Shuffling', 'Disabled (preserve order)')
        ]

        for i, (comp, seed) in enumerate(seed_data, 1):
            table.rows[i].cells[0].text = comp
            table.rows[i].cells[1].text = seed

        self.doc.add_paragraph()

        self.doc.add_heading('7.5 Expected Outputs', level=2)

        self.doc.add_paragraph('After successful execution, the following files should be generated:')

        outputs = [
            'results/experimental_results.json - Raw experimental data',
            'results/exp1_threshold_analysis.png - Experiment 1 visualization',
            'results/exp2_skill_progression.png - Experiment 2 visualization',
            'results/exp3_blockchain_performance.png - Experiment 3 visualization',
            'results/exp4_model_comparison.png - Experiment 4 visualization',
            'results/exp5_scalability.png - Experiment 5 visualization',
            'results/exp6_distribution.png - Experiment 6 visualization',
            'results/similarity_heatmap.png - Similarity matrix heatmap',
            'results/top_recommendations_bar.png - Top recommendations chart',
            'results/recommendations_report.json - Recommendation details',
            'results/verification_report.json - Blockchain verification results',
            'results/SkillChain_DX_Implementation_Results.docx - Research paper'
        ]

        for output in outputs:
            self.doc.add_paragraph(output, style='List Bullet')

        self.doc.add_paragraph()

    def add_validation_criteria(self):
        """Add validation criteria"""
        self.doc.add_heading('8. VALIDATION CRITERIA', level=1)

        self.doc.add_heading('8.1 Data Validation', level=2)

        table = self.doc.add_table(rows=7, cols=3)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Check'
        table.rows[0].cells[1].text = 'Criterion'
        table.rows[0].cells[2].text = 'Pass/Fail'

        data_val = [
            ('Record Counts', 'Roles=15, Courses=25, Employees=10', 'Must match'),
            ('Missing Values', 'No null/empty required fields', 'Zero nulls'),
            ('Skill Format', 'Comma-separated strings', 'Valid format'),
            ('ID Uniqueness', 'All IDs unique within dataset', '100% unique'),
            ('Data Types', 'Correct types for all fields', 'All valid'),
            ('Encoding', 'UTF-8 encoding', 'UTF-8')
        ]

        for i, (check, criterion, pf) in enumerate(data_val, 1):
            table.rows[i].cells[0].text = check
            table.rows[i].cells[1].text = criterion
            table.rows[i].cells[2].text = pf

        self.doc.add_paragraph()

        self.doc.add_heading('8.2 Model Validation', level=2)

        table = self.doc.add_table(rows=6, cols=3)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Check'
        table.rows[0].cells[1].text = 'Criterion'
        table.rows[0].cells[2].text = 'Pass/Fail'

        model_val = [
            ('Model Loading', 'Successfully loads without errors', 'No errors'),
            ('Embedding Dimension', 'Output dimension = 384', '384'),
            ('Score Range', 'All scores in [0, 1]', '0 ≤ score ≤ 1'),
            ('Determinism', 'Same input → same output', '100% match'),
            ('Performance', 'Inference time < 100ms/batch', '<100ms')
        ]

        for i, (check, criterion, pf) in enumerate(model_val, 1):
            table.rows[i].cells[0].text = check
            table.rows[i].cells[1].text = criterion
            table.rows[i].cells[2].text = pf

        self.doc.add_paragraph()

        self.doc.add_heading('8.3 Results Validation', level=2)

        table = self.doc.add_table(rows=8, cols=3)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Check'
        table.rows[0].cells[1].text = 'Criterion'
        table.rows[0].cells[2].text = 'Expected Range'

        results_val = [
            ('Mean Similarity', 'Average score across all pairs', '60-70%'),
            ('Std Deviation', 'Score variability', '10-20%'),
            ('Top Scores', 'Best match per employee', '>75%'),
            ('Coverage', 'Employees with recommendations', '100%'),
            ('Blockchain Throughput', 'Operations per second', '>1000 ops/sec'),
            ('Scalability R²', 'Linear correlation', '>0.95'),
            ('File Generation', 'All output files created', '12 files')
        ]

        for i, (check, criterion, exp_range) in enumerate(results_val, 1):
            table.rows[i].cells[0].text = check
            table.rows[i].cells[1].text = criterion
            table.rows[i].cells[2].text = exp_range

        self.doc.add_paragraph()

    def add_appendix(self):
        """Add appendix with additional information"""
        self.doc.add_heading('9. APPENDIX', level=1)

        self.doc.add_heading('9.1 File Structure', level=2)

        self.doc.add_paragraph('Project directory structure:')

        structure = [
            'test/',
            '├── data/',
            '│   ├── job_roles.csv',
            '│   ├── training_courses.csv',
            '│   └── employees.csv',
            '├── src/',
            '│   ├── skill_matcher.py',
            '│   ├── blockchain.py',
            '│   ├── experiments.py',
            '│   ├── docx_generator.py',
            '│   └── experiment_config_generator.py',
            '├── results/',
            '│   ├── *.png (visualizations)',
            '│   ├── *.json (data files)',
            '│   └── *.docx (documents)',
            '├── run_comprehensive_experiments.py',
            '├── requirements.txt',
            '└── README.md'
        ]

        for line in structure:
            p = self.doc.add_paragraph(line)
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.runs[0]
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

        self.doc.add_paragraph()

        self.doc.add_heading('9.2 Key Dependencies', level=2)

        self.doc.add_paragraph('Complete requirements.txt content:')

        requirements = [
            'sentence-transformers==2.2.2',
            'pandas==2.1.4',
            'numpy==1.26.2',
            'matplotlib==3.8.2',
            'seaborn==0.13.0',
            'scikit-learn==1.3.2',
            'python-docx==1.1.0'
        ]

        for req in requirements:
            p = self.doc.add_paragraph(req)
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.runs[0]
            run.font.name = 'Courier New'

        self.doc.add_paragraph()

        self.doc.add_heading('9.3 Troubleshooting', level=2)

        issues = [
            ('Model download fails', 'Check internet connection; models auto-download on first use'),
            ('Out of memory error', 'Reduce batch size or dataset size; requires 8GB+ RAM'),
            ('Import errors', 'Verify all packages installed: pip install -r requirements.txt'),
            ('Incorrect results', 'Check data file formats; verify Python version ≥3.8'),
            ('Slow performance', 'Normal for first run (model download); subsequent runs faster')
        ]

        for issue, solution in issues:
            p = self.doc.add_paragraph()
            p.add_run(f'Issue: {issue}\n').bold = True
            p.add_run(f'Solution: {solution}')

        self.doc.add_paragraph()

        self.doc.add_heading('9.4 Contact Information', level=2)

        self.doc.add_paragraph(
            'For questions, issues, or collaboration inquiries regarding this experimental '
            'configuration or the SkillChain DX system, please refer to the main research paper '
            'or contact the development team.'
        )

        self.doc.add_paragraph()

        self.doc.add_heading('9.5 Version History', level=2)

        table = self.doc.add_table(rows=4, cols=3)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Version'
        table.rows[0].cells[1].text = 'Date'
        table.rows[0].cells[2].text = 'Changes'

        versions = [
            ('1.0', '2026-01-12', 'Initial release with all 6 experiments'),
            ('1.1', 'TBD', 'Planned: Additional scalability tests'),
            ('2.0', 'TBD', 'Planned: Real-world deployment configuration')
        ]

        for i, (ver, date, changes) in enumerate(versions, 1):
            table.rows[i].cells[0].text = ver
            table.rows[i].cells[1].text = date
            table.rows[i].cells[2].text = changes

        self.doc.add_paragraph()

    def generate(self, output_path='results/SkillChain_DX_Experiment_Configuration.docx'):
        """Generate the complete configuration document"""
        print("\n" + "="*80)
        print("  GENERATING EXPERIMENT CONFIGURATION DOCUMENTATION")
        print("="*80 + "\n")

        print("Adding sections...")
        print("  ✓ Title page")
        self.add_title_page()

        print("  ✓ Section 1: System Configuration")
        self.add_system_configuration()

        print("  ✓ Section 2: Dataset Configuration")
        self.add_dataset_configuration()

        print("  ✓ Section 3: Experiment Configurations")
        self.add_experiment_configurations()

        print("  ✓ Section 4: Algorithm Configurations")
        self.add_algorithm_configurations()

        print("  ✓ Section 5: Evaluation Metrics")
        self.add_evaluation_metrics()

        print("  ✓ Section 6: Visualization Configurations")
        self.add_visualization_configurations()

        print("  ✓ Section 7: Reproducibility Guidelines")
        self.add_reproducibility_guidelines()

        print("  ✓ Section 8: Validation Criteria")
        self.add_validation_criteria()

        print("  ✓ Section 9: Appendix")
        self.add_appendix()

        # Save document
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(output_path)

        print("\n" + "="*80)
        print("  ✓ CONFIGURATION DOCUMENT GENERATED SUCCESSFULLY")
        print(f"  Location: {output_path}")
        print("="*80 + "\n")

        return output_path


def main():
    """Main execution function"""
    generator = ExperimentConfigGenerator()
    output_file = generator.generate()

    print(f"✓ Experiment configuration document ready: {output_file}")
    print("✓ This document contains complete specifications for reproducibility")


if __name__ == '__main__':
    main()

